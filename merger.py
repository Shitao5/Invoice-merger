import os
import re
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from pdf2image import convert_from_path
import tkinter as tk
from tkinter import filedialog, messagebox

def extract_invoice_amount(filename):
    # 使用正则表达式从文件名中提取发票金额
    pattern = r'-(\d+\.\d+)元-'
    match = re.search(pattern, filename)
    if match:
        return match.group(1)
    else:
        return None

def merge_invoices_and_itineraries(folder_path, output_path):
    # 创建一个新的Word文档
    document = Document()

    # 获取文件夹下的所有文件
    files = os.listdir(folder_path)

    # 按照文件名进行排序，以确保发票和行程单的顺序一致
    files.sort()

    # 跟踪发票总金额
    total_amount = 0.0

    # 遍历每个文件
    for file in files:
        file_path = os.path.join(folder_path, file)

        # 检查文件类型是发票还是行程单
        if '发票' in file:
            # 提取发票金额
            invoice_amount = extract_invoice_amount(file)

            # 更新总金额
            if invoice_amount:
                total_amount += float(invoice_amount)

            # 打开PDF文件
            pdf = PdfReader(file_path)

            # 提取第一页作为发票截图
            first_page = pdf.pages[0]

            # 将PDF页面转换为图像
            images = convert_from_path(file_path, dpi=300)

            # 保存图像文件
            invoice_screenshot_path = 'invoice_screenshot.png'
            images[0].save(invoice_screenshot_path, 'PNG')

            # 在Word文档中添加发票金额和截图
            document.add_paragraph('发票金额：{}元'.format(invoice_amount))
            document.add_picture(invoice_screenshot_path, width=Inches(6))

        elif '行程单' in file:
            # 打开PDF文件
            pdf = PdfReader(file_path)

            # 提取第一页作为行程单截图
            first_page = pdf.pages[0]

            # 将PDF页面转换为图像
            images = convert_from_path(file_path, dpi=300)

            # 保存图像文件
            itinerary_screenshot_path = 'itinerary_screenshot.png'
            images[0].save(itinerary_screenshot_path, 'PNG')

            # 在Word文档中添加行程单截图
            document.add_picture(itinerary_screenshot_path, width=Inches(6))

            # 插入分页符
            document.add_page_break()

    # 在Word文档开头添加总金额
    paragraph = document.add_paragraph()
    run = paragraph.add_run('发票总金额：{}元'.format(total_amount))
    font = run.font
    font.size = Pt(14)
    font.bold = True

    # 保存并关闭Word文档
    document.save(output_path)

    # 删除临时保存的截图文件
    os.remove(invoice_screenshot_path)
    os.remove(itinerary_screenshot_path)

    messagebox.showinfo("完成", "合并完成！输出文件路径：{}".format(output_path))

def select_folder():
    folder_path = filedialog.askdirectory()
    folder_path_entry.delete(0, tk.END)
    folder_path_entry.insert(tk.END, folder_path)

def select_output_path():
    output_path = filedialog.asksaveasfilename(defaultextension=".docx")
    output_path_entry.delete(0, tk.END)
    output_path_entry.insert(tk.END, output_path)

def run_merge():
    folder_path = folder_path_entry.get()
    output_path = output_path_entry.get()
    if folder_path and output_path:
        merge_invoices_and_itineraries(folder_path, output_path)
    else:
        messagebox.showerror("错误", "请选择文件夹和输出路径！")

# 创建主窗口
window = tk.Tk()
window.title("发票行程单合并工具")

# 创建选择文件夹的按钮和文本框
folder_label = tk.Label(window, text="选择文件夹:")
folder_label.pack()
folder_path_entry = tk.Entry(window)
folder_path_entry.pack()
folder_button = tk.Button(window, text="选择文件夹", command=select_folder)
folder_button.pack()

# 创建选择输出路径的按钮和文本框
output_label = tk.Label(window, text="选择输出路径:")
output_label.pack()
output_path_entry = tk.Entry(window)
output_path_entry.pack()
output_button = tk.Button(window, text="选择输出路径", command=select_output_path)
output_button.pack()

# 创建运行按钮
run_button = tk.Button(window, text="运行合并", command=run_merge)
run_button.pack()

# 进入主循环
window.mainloop()
